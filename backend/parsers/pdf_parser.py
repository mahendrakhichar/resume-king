"""Resume document parsers for PDF and DOCX formats."""

import re
from pathlib import Path

from utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF with pdfplumber fallback."""
    text = ""

    # Try PyMuPDF first (fastest, best for most PDFs)
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text("text") + "\n"
        doc.close()

        if text.strip():
            logger.info(f"PDF extracted via PyMuPDF: {len(text)} chars")
            return _clean_text(text)
    except Exception as e:
        logger.warning(f"PyMuPDF failed for {file_path}: {e}")

    # Fallback to pdfplumber (better for complex layouts)
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        if text.strip():
            logger.info(f"PDF extracted via pdfplumber: {len(text)} chars")
            return _clean_text(text)
    except Exception as e:
        logger.error(f"pdfplumber also failed for {file_path}: {e}")

    raise ValueError(f"Could not extract text from PDF: {file_path}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n".join(paragraphs)
        logger.info(f"DOCX extracted: {len(text)} chars")
        return _clean_text(text)

    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        raise ValueError(f"Could not extract text from DOCX: {file_path}")


def extract_text(file_path: str) -> str:
    """Extract text from a resume file based on its extension."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _clean_text(text: str) -> str:
    """Clean extracted text: normalize whitespace, remove artifacts."""
    # Replace multiple newlines with double newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Replace multiple spaces with single space
    text = re.sub(r" {2,}", " ", text)
    # Remove null bytes and control characters (except newline/tab)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return text.strip()
